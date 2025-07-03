# cx_vendor_interview_app.py

import streamlit as st
import openai
from openai import OpenAI
import json
import os
from docx import Document
from datetime import datetime
from dotenv import load_dotenv

#load environment variables from .env file
load_dotenv()

# --- CONFIGURATION ---
#set API key from environment variable
##openai.api_key = os.getenv("OPENAI_API_KEY")

##new OpenAI client object
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

docx_filename = "vendor_responses.docx"
json_filename = "vendor_responses.json"

def load_questions():

    # Simplified version of all sections extracted from your Word document.
    return [
        "When was your platform founded?",
        "Who was the founder and what is their current position in your company?",
        "In a nutshell, what does your platform do?",
        "At a high level what are the key features of your platform?",
        "What size businesses do you support?",
        "What sectors do you have specialism in / have most of your clients?",
        "Can you do B2B?",
        "Please provide a case study. This should highlight what makes your platform special.",
        "How does your platform upload historical feedback?",
        "What channels can clients use to distribute surveys?",
        "Which review platforms can your platform access?",
        "What formats of customer support data (e.g. chat, calls, emails) can your platform access?",
        "Can your platform identify churn probability by customer?",
        "Can your platform highlight touchpoints with significant issues?",
        "How does your platform help clients close individual issues?",
        "What sets your platform apart from other VoC solutions?",
        "If your platform were a person, who would they be and why?"
    ]

def ask_gpt_followup(question, answer):
    prompt = f"""
    A vendor answered the following survey question:
    Question: {question}
    Answer: {answer}
    
    Is this answer vague, incomplete, or too short? If yes, suggest a follow-up question.
    If it is sufficient, respond only with: 'Good answer.'
    """
    try:
        response = client.chat.completions.create(
            model="gpt-4-turbo",
            messages=[{"role": "user", "content": prompt}]
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        return f"Error: {e}"

def save_to_docx(answers):
    doc = Document()
    doc.add_heading("CX Vendor Survey Responses", 0)
    doc.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
    for q, a in answers.items():
        doc.add_heading(q, level=2)
        doc.add_paragraph(a)
    doc.save(docx_filename)

def save_to_json(answers):
    with open(json_filename, "w") as f:
        json.dump(answers, f, indent=2)

# --- STREAMLIT APP ---
st.set_page_config(page_title="CX Vendor Interview Agent", layout="wide")
st.title("🧠 CX Vendor Interview Agent")

questions = load_questions()

if "answers" not in st.session_state:
    st.session_state.answers = {}

for idx, question in enumerate(questions):
    if question in st.session_state.answers:
        continue

    st.subheader(f"Question {idx + 1} of {len(questions)}")
    st.markdown(f"**{question}**")
    user_input = st.text_area("Your answer:", key=f"input_{idx}")

    if st.button("Submit", key=f"submit_{idx}"):
        feedback = ask_gpt_followup(question, user_input)
        if "Good answer." in feedback:
            st.session_state.answers[question] = user_input
            st.success("✅ Answer accepted.")
        else:
            st.warning(f"⚠️ AI Feedback: {feedback}")
        st.stop()

if len(st.session_state.answers) == len(questions):
    st.success("🎉 All responses completed!")
    save_to_json(st.session_state.answers)
    save_to_docx(st.session_state.answers)
    st.download_button("📄 Download Word Doc", data=open(docx_filename, "rb"), file_name=docx_filename)
    st.download_button("📄 Download JSON", data=open(json_filename, "rb"), file_name=json_filename)
